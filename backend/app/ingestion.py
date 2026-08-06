"""
Document ingestion pipeline — text extraction and semantic-aware chunking.
Supports PDF, DOCX, TXT, CSV, and Markdown files.
"""
import os
import re
import logging
from typing import List, Dict, Any

logger = logging.getLogger("processpilot.ingestion")

# Optional imports with safe fallbacks
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import docx
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

def extract_text_from_pdf(file_path: str) -> str:
    if not HAS_PYMUPDF:
        return "[PyMuPDF not installed] Fallback: Please install pymupdf to parse PDF text."
    try:
        text = ""
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text() + "\n"
        return text
    except Exception as e:
        return f"[PDF parsing error]: {str(e)}"

def extract_text_from_docx(file_path: str) -> str:
    if not HAS_PYTHON_DOCX:
        return "[python-docx not installed] Fallback: Please install python-docx to parse Word documents."
    try:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        return f"[Word parsing error]: {str(e)}"

def extract_text_from_excel(file_path: str) -> str:
    if not HAS_PANDAS:
        return "[pandas not installed] Fallback: Please install pandas to parse Excel documents."
    try:
        text_parts = []
        df_dict = pd.read_excel(file_path, sheet_name=None)
        for sheet_name, df in df_dict.items():
            text_parts.append(f"## Sheet: {sheet_name}")
            text_parts.append(df.to_csv(index=False))
            text_parts.append("\n")
        return "\n".join(text_parts)
    except Exception as e:
        return f"[Excel parsing error]: {str(e)}"

def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        return f"[Text file reading error]: {str(e)}"

def extract_content(file_path: str, file_type: str) -> str:
    """Extract text from a document and automatically redact PII before returning."""
    from .pii_redactor import redact_document  # lazy import to avoid circular deps
    ext = file_type.lower()
    if ext == "pdf":
        raw_text = extract_text_from_pdf(file_path)
    elif ext in ["docx", "doc"]:
        raw_text = extract_text_from_docx(file_path)
    elif ext in ["xlsx", "xls"]:
        raw_text = extract_text_from_excel(file_path)
    else:
        raw_text = extract_text_from_txt(file_path)

    # PII Redaction Gate — strip PII before any data hits ChromaDB or LLM
    redacted_text, count = redact_document(raw_text)
    if count > 0:
        logger.info(f"[Ingestion] PII redacted {count} entities from '{file_path}'")
    return redacted_text


def chunk_text(text: str, chunk_size: int = 800, chunk_overlap: int = 150) -> List[str]:
    """
    Semantic-aware recursive text splitter.
    Splits by hierarchy: sections → paragraphs → sentences → words.
    Preserves structural formatting better than naive character splitting.
    """
    if not text or not text.strip():
        return []
    
    # Normalize excessive whitespace while preserving paragraph breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = text.strip()
    
    if len(text) <= chunk_size:
        return [text]
    
    # Separators in order of priority (most meaningful boundaries first)
    separators = ["\n## ", "\n### ", "\n\n", "\n", ". ", "; ", ", ", " "]
    
    def _split_recursive(text: str, seps: List[str]) -> List[str]:
        if len(text) <= chunk_size:
            return [text.strip()] if text.strip() else []
        
        # Try each separator starting from most meaningful
        for sep in seps:
            if sep in text:
                parts = text.split(sep)
                chunks = []
                current = ""
                
                for part in parts:
                    candidate = current + sep + part if current else part
                    
                    if len(candidate) <= chunk_size:
                        current = candidate
                    else:
                        if current.strip():
                            chunks.append(current.strip())
                        
                        # If single part exceeds chunk_size, recurse with finer separators
                        if len(part) > chunk_size:
                            remaining_seps = seps[seps.index(sep) + 1:]
                            if remaining_seps:
                                chunks.extend(_split_recursive(part, remaining_seps))
                            else:
                                # Last resort: hard split with overlap
                                for i in range(0, len(part), chunk_size - chunk_overlap):
                                    chunk = part[i:i + chunk_size]
                                    if chunk.strip():
                                        chunks.append(chunk.strip())
                        else:
                            current = part
                
                if current.strip():
                    chunks.append(current.strip())
                
                if chunks:
                    return chunks
        
        # No separator found — hard split with overlap
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start += chunk_size - chunk_overlap
        return chunks
    
    return _split_recursive(text, separators)


def process_file_upload(file_path: str, file_type: str, document_id: int) -> List[Dict[str, Any]]:
    """
    Extracts text, chunks it, and returns the chunks ready for vector database insertion.
    Reads from a local file path — used in legacy / local-dev code paths.
    """
    raw_text = extract_content(file_path, file_type)
    text_chunks = chunk_text(raw_text)
    
    processed_chunks = []
    for idx, chunk in enumerate(text_chunks):
        processed_chunks.append({
            "id": f"doc_{document_id}_chunk_{idx}",
            "text": chunk,
            "index": idx,
            "metadata": {
                "file_name": os.path.basename(file_path),
                "file_type": file_type
            }
        })
    return processed_chunks


def extract_content_from_bytes(file_bytes: bytes, file_type: str, filename: str = "document") -> str:
    """
    Extract and PII-redact text directly from in-memory bytes.
    Used by background tasks that receive file content after the HTTP request
    has already returned (Supabase Storage path or in-memory bytes from upload).
    """
    from .pii_redactor import redact_document
    import io
    ext = file_type.lower()
    raw_text = ""

    if ext == "pdf":
        if not HAS_PYMUPDF:
            return "[PyMuPDF not installed] Cannot parse PDF from bytes."
        try:
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                for page in doc:
                    raw_text += page.get_text() + "\n"
        except Exception as e:
            return f"[PDF bytes parsing error]: {str(e)}"

    elif ext in ["docx", "doc"]:
        if not HAS_PYTHON_DOCX:
            return "[python-docx not installed] Cannot parse DOCX from bytes."
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            raw_text = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            return f"[DOCX bytes parsing error]: {str(e)}"

    elif ext in ["xlsx", "xls"]:
        if not HAS_PANDAS:
            return "[pandas not installed] Cannot parse Excel from bytes."
        try:
            import io
            df_dict = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
            text_parts = []
            for sheet_name, df in df_dict.items():
                text_parts.append(f"## Sheet: {sheet_name}")
                text_parts.append(df.to_csv(index=False))
                text_parts.append("")
            raw_text = "\n".join(text_parts)
        except Exception as e:
            return f"[Excel bytes parsing error]: {str(e)}"

    else:
        # TXT, CSV, MD — decode as UTF-8
        try:
            raw_text = file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            return f"[Text decode error]: {str(e)}"

    # PII Redaction Gate
    redacted_text, count = redact_document(raw_text)
    if count > 0:
        logger.info(f"[Ingestion] PII redacted {count} entities from '{filename}'")
    return redacted_text


def process_file_upload_from_bytes(
    file_bytes: bytes,
    file_type: str,
    document_id: int,
    filename: str = "document",
) -> List[Dict[str, Any]]:
    """
    Bytes-native ingestion pipeline — the canonical path used by BackgroundTasks.
    No disk I/O required: operates entirely on the in-memory bytes passed from
    the upload handler, eliminating any dependency on the local filesystem.
    """
    raw_text = extract_content_from_bytes(file_bytes, file_type, filename)
    text_chunks = chunk_text(raw_text)

    processed_chunks = []
    for idx, chunk in enumerate(text_chunks):
        processed_chunks.append({
            "id": f"doc_{document_id}_chunk_{idx}",
            "text": chunk,
            "index": idx,
            "metadata": {
                "file_name": filename,
                "file_type": file_type,
            },
        })
    return processed_chunks
